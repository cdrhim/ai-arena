from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE = Path(r"C:\Users\임해룡\Documents\SparkLabs AI Arena")
OUTPUT_DIR = WORKSPACE / "deliverables"
OUTPUT_PATH = OUTPUT_DIR / "SparkClaw_AI_Arena_제품_방향_및_핵심_스펙_v0.1.docx"
SKILL_SCRIPTS = Path(
    r"C:\Users\임해룡\.codex\plugins\cache\openai-primary-runtime\documents\26.812.11052\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


NAVY = "0B2348"
DEEP_BLUE = "123A73"
BLUE = "1769E8"
TEAL = "2FCDB3"
INK = "13233F"
MUTED = "66758D"
PALE_BLUE = "EAF2FF"
PALE_TEAL = "E9F9F6"
PALE_GRAY = "F3F6FA"
LINE = "D6E1F0"
WHITE = "FFFFFF"
GOLD = "F2B84B"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    latin: str = "Calibri",
    east_asia: str = "맑은 고딕",
) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_font(paragraph, **kwargs) -> None:
    for run in paragraph.runs:
        set_run_font(run, **kwargs)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_borders(paragraph, *, accent: str, line: str = LINE) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge, color, size in (
        ("top", line, 8),
        ("bottom", line, 8),
        ("left", accent, 24),
        ("right", line, 8),
    ):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "5")
        element.set(qn("w:color"), color)


def set_keep_with_next(paragraph, keep: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    element = p_pr.find(qn("w:keepNext"))
    if keep and element is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not keep and element is not None:
        p_pr.remove(element)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def add_text(
    doc: Document,
    text: str,
    *,
    size: float = 11,
    color: str = INK,
    bold: bool = False,
    italic: bool = False,
    before: float = 0,
    after: float = 6,
    line: float = 1.10,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    keep_next: bool = False,
) -> object:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    set_keep_with_next(p, keep_next)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> object:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_keep_with_next(p, True)
    return p


def add_bullet(doc: Document, text: str, *, bold_prefix: str | None = None) -> object:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)
    return p


def add_numbered(doc: Document, text: str) -> object:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_callout(
    doc: Document,
    label: str,
    body: str,
    *,
    fill: str = PALE_BLUE,
    accent: str = BLUE,
) -> object:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    set_paragraph_shading(p, fill)
    set_paragraph_borders(p, accent=accent)
    r = p.add_run(label.upper())
    set_run_font(r, size=9.5, color=accent, bold=True)
    r.add_break()
    r2 = p.add_run(body)
    set_run_font(r2, size=11.5, color=INK, bold=True)
    return p


def format_table_text(cell, *, bold: bool = False, color: str = INK, size: float = 10.2) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.10
        for run in p.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_lifecycle_table(doc: Document) -> None:
    rows = [
        ("1", "탐색", "기업·역량·Task를 공개 근거로 파악"),
        ("2", "대화", "Community에서 질문·경험·공지를 축적"),
        ("3", "검토 요청", "협업 목적과 메시지를 상대 팀에 전달"),
        ("4", "상대 동의", "상대 팀이 My Log에서 승인 또는 거절"),
        ("5", "연결", "승인 후 SparkLabs가 다음 단계와 소개를 조율"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0].cells
    header[0].text = "단계"
    header[1].text = "사용자 행동"
    header[2].text = "Arena의 역할"
    for cell in header:
        set_cell_shading(cell, NAVY)
        format_table_text(cell, bold=True, color=WHITE, size=10)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for number, action, role in rows:
        cells = table.add_row().cells
        cells[0].text = number
        cells[1].text = action
        cells[2].text = role
        for idx, cell in enumerate(cells):
            set_cell_shading(cell, WHITE if int(number) % 2 else PALE_GRAY)
            format_table_text(cell, bold=(idx == 1), size=10.2)
            if idx == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 6, "color": LINE},
                bottom={"val": "single", "sz": 6, "color": LINE},
                left={"val": "single", "sz": 6, "color": LINE},
                right={"val": "single", "sz": 6, "color": LINE},
            )
    apply_table_geometry(table, [900, 2100, 6360], indent_dxa=120)


def add_role_block(doc: Document, title: str, purpose: str, needs: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.10
    set_keep_with_next(p, True)
    r = p.add_run(title)
    set_run_font(r, size=12, color=DEEP_BLUE, bold=True)
    r2 = p.add_run(f"  {purpose}")
    set_run_font(r2, size=10.5, color=MUTED, italic=True)
    for item in needs:
        add_bullet(doc, item)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DEEP_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = p.add_run("SPARKCLAW AI ARENA  |  CONTEXT BRIEF")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = p.add_run("\tV0.1  |  2026.08")
    set_run_font(right, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    prefix = fp.add_run("SparkClaw AI Arena  ·  내부 협업용  ·  ")
    set_run_font(prefix, size=8.5, color=MUTED)
    add_page_field(fp)


def build() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        configure_header_footer(section)

    # Cover / memo masthead
    kicker = add_text(
        doc,
        "PRODUCT CONTEXT BRIEF",
        size=10,
        color=TEAL,
        bold=True,
        before=16,
        after=5,
        keep_next=True,
    )
    set_paragraph_shading(kicker, NAVY)
    kicker.paragraph_format.left_indent = Inches(0.18)
    kicker.paragraph_format.right_indent = Inches(0.18)
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(0)
    kicker.paragraph_format.line_spacing = 1.0

    title = add_text(
        doc,
        "SparkClaw AI Arena\n제품 방향 및 핵심 스펙",
        size=28,
        color=WHITE,
        bold=True,
        before=0,
        after=0,
        line=1.08,
        keep_next=True,
    )
    set_paragraph_shading(title, NAVY)
    title.paragraph_format.left_indent = Inches(0.18)
    title.paragraph_format.right_indent = Inches(0.18)

    subtitle = add_text(
        doc,
        "Purpose, Audience & Core Product Scope",
        size=13,
        color="BFD6FF",
        before=0,
        after=18,
        line=1.0,
    )
    set_paragraph_shading(subtitle, NAVY)
    subtitle.paragraph_format.left_indent = Inches(0.18)
    subtitle.paragraph_format.right_indent = Inches(0.18)

    add_text(
        doc,
        "호민 대표님과의 피드백 및 수정 논의, 현재 구현 방향을 바탕으로 정리한 협업용 컨텍스트 문서",
        size=10.5,
        color=MUTED,
        bold=True,
        after=10,
    )

    add_callout(
        doc,
        "한 문장 정의",
        "SparkClaw AI Arena는 검증된 AI 스타트업, 창업자, 산업 파트너와 SparkLabs 네트워크가 서로를 발견하고 대화하며, 상대의 동의를 거쳐 실제 협업으로 이어지는 AI-native 커뮤니티 플랫폼입니다.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    add_text(doc, "문서 한눈에 보기", size=12, color=DEEP_BLUE, bold=True, before=8, after=5)
    overview = [
        "왜 만드는가: 프로그램 운영 정보와 외부 연결 경험을 분리하고, Arena를 실제 만남과 협업의 접점으로 만든다.",
        "누구를 위한가: 참여 스타트업, 산업 파트너, 멘토·투자자·LP, SparkLabs 운영진.",
        "무엇이 핵심인가: Discover, Community, 동의 기반 협업 요청, Clawee, My Log.",
        "무엇을 나중에 여는가: 실제 스폰서가 승인된 Bounty, 외부 알림과 더 넓은 생태계 연동.",
    ]
    for item in overview:
        add_bullet(doc, item)

    note = add_text(
        doc,
        "문서 상태  ·  v0.1  ·  2026년 8월  ·  내부 협업 및 온보딩용",
        size=9.5,
        color=MUTED,
        italic=True,
        before=10,
        after=0,
    )
    note.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_page_break()

    add_heading(doc, "1. 문서 목적", 1)
    add_text(
        doc,
        "이 문서는 중간에 합류한 협업자가 SparkClaw AI Arena의 전후 맥락을 빠르게 파악하고, 이후 제품·운영·데이터·콘텐츠 의사결정을 같은 기준에서 진행할 수 있도록 만든 기준 문서입니다.",
    )
    for item in [
        "현재까지 어떤 문제를 해결하려 했는지",
        "Arena가 기존 SparkClaw 프로그램 사이트와 어떻게 다른지",
        "사용자별로 무엇을 제공해야 하는지",
        "현재 반드시 포함할 기능과 준비 단계로 남겨둘 기능은 무엇인지",
        "운영 시 지켜야 할 동의·검증·개인정보 원칙은 무엇인지",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. 추진 배경", 1)
    add_text(
        doc,
        "SparkClaw에는 다양한 AI 스타트업이 참여하지만, 각 팀이 다른 팀의 제품명만 보고 실제로 어떤 문제를 해결하는지, 어떤 역량과 검증 근거가 있는지, 지금 누구와 연결되면 좋은지 파악하기는 어렵습니다. 기존 프로그램 운영 정보만으로는 '함께할 이유'가 충분히 드러나지 않습니다.",
    )
    add_text(doc, "미팅에서 반복적으로 확인된 수요", size=12, color=DEEP_BLUE, bold=True, before=8, after=4)
    for item in [
        "회계·재무·법무·채용 등 창업 운영 문제를 서로 묻고 실제 경험을 나누고 싶다.",
        "개발 중심 팀은 마케팅, 고객 획득, 영업, PoC 설계 경험이 필요하다.",
        "비개발 창업자는 개발 외주, 위탁, 기술 파트너 선정 경험이 필요하다.",
        "상대 연락처를 바로 공개하지 않고도 안전하게 소개를 요청하고 싶다.",
        "산업 파트너는 전체 참가기업을 보고, 자신의 과제와 맞는 기업을 근거 중심으로 찾고 싶다.",
        "SparkLabs 운영진은 추천·소개·공지·활동 이력을 확인하고 운영할 수 있어야 한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. 기존 사이트와 Arena의 역할 구분", 1)
    add_callout(
        doc,
        "역할 구분",
        "기존 SparkClaw 프로그램 사이트는 프로그램 참여와 운영을 위한 내부 허브이고, AI Arena는 참여팀·파트너·SparkLabs 네트워크가 서로를 발견하고 연결되는 외부 지향 관계 레이어입니다.",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_bullet(doc, "프로그램 사이트: Weekly Report, 일정, 혜택, 과제, 운영 공지, 프로그램 내부 데이터.")
    add_bullet(doc, "AI Arena: 기업 탐색, Task 기반 검색, Community, 협업 검토 요청, 동의 기반 소개, 개인 활동 로그.")
    add_text(
        doc,
        "따라서 Arena에서 프로그램 사이트의 모든 기능을 다시 만들기보다, '누구를 만나고 무엇을 함께할 것인가'에 집중해야 합니다.",
        bold=True,
        color=DEEP_BLUE,
        before=4,
    )

    add_heading(doc, "4. 핵심 사용자와 기대 가치", 1)
    add_role_block(
        doc,
        "SparkClaw 참여기업",
        "동료 팀과 실제 문제를 나누고 협업 기회를 찾는 사용자",
        [
            "다른 팀의 서비스, 핵심 Task, 팀 경쟁력, 정량 근거를 빠르게 이해한다.",
            "나와 비슷하거나 상호보완적인 팀을 발견하고 비교한다.",
            "Community에서 질문·경험·출시 소식·구하는 도움을 게시하고 반응을 얻는다.",
            "협업 요청과 응답, 댓글과 Bounty 진행 상황을 My Log에서 확인한다.",
        ],
    )
    add_role_block(
        doc,
        "산업 파트너·외부 기업",
        "자사의 문제를 해결할 스타트업을 탐색하는 사용자",
        [
            "산업명보다 해결하려는 Task와 검증 근거를 기준으로 기업을 찾는다.",
            "공개 가능한 프로필만 비교하고, 소개는 상대 팀 동의 후 진행한다.",
            "PoC·평가·API 파트너·채용·사업 제휴 목적을 명확히 전달한다.",
        ],
    )
    add_role_block(
        doc,
        "멘토·투자자·LP",
        "팀의 실행력과 성장 가능성을 판단하는 네트워크 사용자",
        [
            "팀 구성, 경력·학력, 해결 문제, 최근 성과와 숫자를 함께 본다.",
            "홍보 문구보다 공개 근거, 검증 단계, 현재 부족한 요소를 확인한다.",
            "멘토링·투자 검토·산업 연결이 필요한 팀을 선별한다.",
        ],
    )
    add_role_block(
        doc,
        "SparkLabs 운영진",
        "큐레이션·소개·운영 품질을 책임지는 관리자",
        [
            "기업 프로필, Spotlight, 공식 공지와 추천 근거를 편집한다.",
            "협업 요청의 발신·수신·승인·거절 이력을 감사한다.",
            "비공개 연락처와 내부 메모를 보호하고 AI의 미확인 추정을 방지한다.",
            "실제 스폰서 Brief가 준비된 Bounty만 승인해 공개한다.",
        ],
    )

    doc.add_page_break()

    add_heading(doc, "5. 제품의 핵심 경험", 1)
    add_heading(doc, "5.1 Discover - 이름이 아니라 해결 역량을 찾는 공간", 2)
    add_text(
        doc,
        "Discover는 단순 기업 목록이 아니라, 지금 이 팀이 무엇을 해결하며 왜 주목할 가치가 있는지 이해시키는 첫 화면입니다.",
    )
    for item in [
        "4 Picks Editorial Spotlight: Weekly Report와 최근 확인된 실행 성과를 바탕으로 운영진이 큐레이션한 기업을 소개한다.",
        "Spark AI 기업 탐색: 사용자가 필요한 도움이나 문제를 자연어로 입력하면 공개 프로필 근거에 맞는 후보를 정리한다.",
        "Task-driven Search: 회사명·산업보다 구체적인 해결 Task를 중심으로 전체 역량을 검색한다.",
        "Company Directory: 전체 참가기업의 안전한 공개 필드를 탐색한다.",
        "Compare: 동일한 질문과 근거 차원으로 최대 3개 기업을 비교한다.",
        "협업 적합 기업·나와 비슷한 팀: 점수 대신 추천 순위와 기업별로 다른 선정 이유를 보여준다.",
    ]:
        add_bullet(doc, item)

    add_callout(
        doc,
        "기업 카드 원칙",
        "서비스 설명만이 아니라 팀 경쟁력, 최근 성과, 정량적 검증, 현 단계, 가장 잘 해결하는 Task와 추가로 해결 가능한 Task를 근거와 함께 보여줍니다.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    add_heading(doc, "5.2 Community - 프로그램 사이의 대화를 축적하는 공간", 2)
    add_text(
        doc,
        "Community는 빈 포럼이 아니라 운영진 공지와 실제 수요가 먼저 보이고, 구성원이 자신의 질문과 경험을 이어서 쌓는 공간입니다.",
    )
    for item in [
        "Start a Conversation: 사용자가 내용을 먼저 쓰고 Spark AI가 제목·채널·공개 범위를 제안한다.",
        "Pre-OT Networking Needs: 회계·재무 운영, 개발팀의 고객 획득, 비개발 외주·위탁 등 실제 공통 수요를 질문 템플릿으로 제공한다.",
        "Community Feed: 인기순, 최신순, 댓글 필요 필터를 제공하고 upvote·댓글·수정이 가능하다.",
        "공개 범위: Public은 SparkClaw 산업 파트너 포함, Private은 SparkClaw 부트캠프 멤버와 SparkLabs만 접근한다.",
        "공식 공지: SparkLabs 운영진이 작성한 AI Arena 전용 공지가 Discover의 What's New와 Community에 함께 반영된다.",
        "유연한 채널: 사용 중인 채널을 우선 노출하고, 사용자가 새 채널을 직접 만들 수 있다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.3 동의 기반 협업 요청", 2)
    add_text(
        doc,
        "소개는 '추천 카드 클릭 → 연락처 공개'가 아니라 상대 팀의 의사를 먼저 확인하는 절차입니다. 이 흐름이 Arena의 신뢰를 만듭니다.",
    )
    add_lifecycle_table(doc)
    add_text(
        doc,
        "연락처는 자동으로 공개하지 않으며, 공개 프로필에서 확인되지 않은 고객·성과·기술 정보는 소개 전에 검증합니다.",
        color=MUTED,
        italic=True,
        before=5,
    )

    add_heading(doc, "5.4 Clawee 클로이 - 안내를 넘어 행동을 연결하는 인터페이스", 2)
    for item in [
        "사이트 목적과 사용법을 설명하고 실제 메뉴·섹션을 이동하며 튜토리얼한다.",
        "기업 찾기, 글 작성, 비교, Bounty와 My Log 사용법을 자연어로 안내한다.",
        "사용자의 명령을 해석해 적합한 화면과 다음 행동으로 연결한다.",
        "장기적으로는 사용자의 이전 탐색·요청·관심사를 기억해 맞춤 제안을 이어간다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.5 My Log - 내 활동과 다음 행동을 보는 기록", 2)
    for item in [
        "Discover에서 보낸 협업 검토 요청과 상대 팀의 응답",
        "Community에서 작성한 글·댓글·반응과 내 글에 들어온 댓글·upvote",
        "Bounty 신청·제출·상태 변경과 다음 단계",
        "모든 활동을 최신순 raw log 형태로 조회하고 필요한 영역으로 돌아가는 링크",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "6. Bounty의 현재 위치", 1)
    add_callout(
        doc,
        "현재 원칙",
        "Bounty는 Arena의 중요한 확장 기능이지만, 실제 스폰서 문제와 운영 승인 없이 데모 과제를 공개하지 않습니다. 준비 단계에서는 취지와 Release Gate만 보여줍니다.",
        fill="FFF7E6",
        accent=GOLD,
    )
    for step in [
        "파트너·기업이 Sponsor Brief로 실제 문제, 성공 기준, 데이터·보안 제약을 제출한다.",
        "SparkLabs가 공개 범위, 평가 기준, 책임자, 일정과 보상을 검토한다.",
        "승인된 Brief만 Bounty Board에 공개하고 참가기업의 신청을 받는다.",
        "제출 결과는 공개·비공개 자동 검증과 전문가 비교 평가를 거친다.",
        "검증된 결과를 PoC, 투자, 크레딧, 사업 연결 같은 실제 기회로 이어간다.",
    ]:
        add_numbered(doc, step)

    add_heading(doc, "7. 제품·운영 원칙", 1)
    principles = [
        ("Community first", "기능을 많이 보여주는 것보다 대화와 관계가 반복해서 생기는 구조를 우선합니다."),
        ("운영 허브와 연결 레이어 분리", "일정·혜택·리포트는 프로그램 사이트, 탐색·대화·소개·활동 기록은 Arena에 둡니다."),
        ("근거 중심 매력도", "기업 카드는 투자자·파트너가 궁금해할 팀 경쟁력, 숫자, 성과, Task와 출처를 중심으로 구성합니다."),
        ("AI는 요약하고 사람은 검증", "AI가 추천·비교·게시 설정을 지원하지만, 미확인 정보를 사실처럼 만들지 않습니다."),
        ("동의와 개인정보 보호", "연락처와 내부 정보는 비공개로 유지하고, 소개는 상대 팀 승인 후 진행합니다."),
        ("경쟁보다 연결", "점수와 서열만 강조하지 않고, 서로 다른 강점과 상호보완 가능성을 보여줍니다."),
    ]
    for label, body in principles:
        add_bullet(doc, f"{label}: {body}", bold_prefix=f"{label}:")

    add_heading(doc, "8. 핵심 스코프와 단계별 확장", 1)
    add_text(doc, "현재 반드시 포함해야 할 범위", size=12, color=DEEP_BLUE, bold=True, before=6, after=4)
    for item in [
        "기업 프로필과 Task 기반 Directory/Search",
        "Editorial Spotlight와 What's New 공식 공지",
        "기업 비교, 나와 비슷한 팀, 협업 적합 기업 추천",
        "Community 게시글·댓글·upvote·수정·공개 범위",
        "상대 팀 동의 기반 협업 검토 요청과 운영진 감사 로그",
        "My Log와 Supabase sc_arena_ 활동 원장",
        "역할별 권한과 공개 범위",
        "Clawee 클로이 튜토리얼과 AI Arena 안내",
    ]:
        add_bullet(doc, item)

    add_text(doc, "다음 단계", size=12, color=DEEP_BLUE, bold=True, before=8, after=4)
    for item in [
        "실제 승인된 Sponsor Brief 기반 Bounty 공개",
        "협업 요청·댓글·Bounty 상태·공식 공지에 대한 인앱 및 이메일 알림",
        "멘토·투자사·LP 네트워크와 팀 데이터의 동의 기반 연결",
        "Clawee의 사용자 맥락·활동 기억과 후속 제안",
        "코호트와 프로그램을 넘어선 SparkLabs 생태계 확장",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9. 성공 기준", 1)
    add_callout(
        doc,
        "최종 판단 질문",
        "SparkClaw에 들어오면 혼자서는 만나기 어려운 사람, 기업, 자원과 기회를 실제로 만날 수 있는가?",
        fill=PALE_TEAL,
        accent=TEAL,
    )
    add_text(
        doc,
        "AI Arena의 성공은 페이지 수나 추천 점수의 정교함이 아니라, 발견한 팀이 대화하고 서로 동의한 뒤 실제 협업·PoC·멘토링·투자 검토로 이어지는지로 판단해야 합니다.",
        size=11.5,
        bold=True,
        color=DEEP_BLUE,
    )

    add_heading(doc, "부록. 정리 근거", 1)
    sources = [
        "260806 1700 AI Arena 호민대표님 피드백.txt",
        "260811 1500 SparkClaw AI Arena 수정논의 미팅.txt",
        "260812 1030 AI Arena 호민대표님.txt",
        "2026년 8월 기준 SparkClaw AI Arena 구현 및 운영 논의 기록",
    ]
    for source in sources:
        add_bullet(doc, source)
    add_text(
        doc,
        "주의: 본 문서는 현재까지 합의된 제품 방향을 정리한 컨텍스트 브리프입니다. 세부 운영 정책, 외부 연동, Bounty 공개 일정은 실제 운영 준비와 데이터 검증 상태에 따라 갱신됩니다.",
        size=9.5,
        color=MUTED,
        italic=True,
        before=8,
    )

    doc.core_properties.title = "SparkClaw AI Arena 제품 방향 및 핵심 스펙"
    doc.core_properties.subject = "Purpose, Audience & Core Product Scope"
    doc.core_properties.author = "SparkLabs"
    doc.core_properties.keywords = "SparkClaw, AI Arena, product context, community, discover, collaboration"

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build())
