from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_ai_arena_context_brief import (
    BLUE,
    DEEP_BLUE,
    GOLD,
    INK,
    MUTED,
    NAVY,
    PALE_BLUE,
    PALE_TEAL,
    TEAL,
    WHITE,
    add_bullet,
    add_callout,
    add_heading,
    add_lifecycle_table,
    add_page_field,
    add_text,
    configure_header_footer,
    configure_section,
    set_paragraph_shading,
    set_run_font,
)


WORKSPACE = Path(r"C:\Users\임해룡\Documents\SparkLabs AI Arena")
OUTPUT_PATH = (
    WORKSPACE
    / "deliverables"
    / "SparkClaw_AI_Arena_제품_방향_및_핵심_스펙_v0.1_4쪽_압축본.docx"
)


def configure_compact_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.18

    tokens = {
        "Heading 1": (16, BLUE, 12, 7),
        "Heading 2": (13, BLUE, 8, 4),
        "Heading 3": (12, DEEP_BLUE, 6, 3),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18


def compact_bullet(doc: Document, text: str, *, bold_prefix: str | None = None):
    p = add_bullet(doc, text, bold_prefix=bold_prefix)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    for run in p.runs:
        if run.bold:
            set_run_font(run, size=10.5, color=INK, bold=True)
        else:
            set_run_font(run, size=10.5, color=INK)
    return p


def role(doc: Document, title: str, intent: str, items: list[str]) -> None:
    p = add_text(
        doc,
        title,
        size=11.5,
        color=DEEP_BLUE,
        bold=True,
        before=5,
        after=1,
        keep_next=True,
    )
    r = p.add_run(f"  {intent}")
    set_run_font(r, size=9.5, color=MUTED, italic=True)
    for item in items:
        compact_bullet(doc, item)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build() -> Path:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_compact_styles(doc)
    for section in doc.sections:
        configure_section(section)
        configure_header_footer(section)

    # PAGE 1 - purpose and context
    kicker = add_text(
        doc,
        "PRODUCT CONTEXT BRIEF  ·  4-PAGE EDITION",
        size=9.5,
        color=TEAL,
        bold=True,
        before=10,
        after=0,
        keep_next=True,
    )
    set_paragraph_shading(kicker, NAVY)
    kicker.paragraph_format.left_indent = Inches(0.16)
    kicker.paragraph_format.right_indent = Inches(0.16)

    title = add_text(
        doc,
        "SparkClaw AI Arena\n제품 방향 및 핵심 스펙",
        size=25,
        color=WHITE,
        bold=True,
        before=0,
        after=0,
        line=1.05,
        keep_next=True,
    )
    set_paragraph_shading(title, NAVY)
    title.paragraph_format.left_indent = Inches(0.16)
    title.paragraph_format.right_indent = Inches(0.16)

    subtitle = add_text(
        doc,
        "Purpose, Audience & Core Product Scope  |  v0.1  |  2026.08",
        size=10.5,
        color="BFD6FF",
        before=0,
        after=9,
    )
    set_paragraph_shading(subtitle, NAVY)
    subtitle.paragraph_format.left_indent = Inches(0.16)
    subtitle.paragraph_format.right_indent = Inches(0.16)

    add_callout(
        doc,
        "한 문장 정의",
        "검증된 AI 스타트업, 창업자, 산업 파트너와 SparkLabs 네트워크가 서로를 발견하고 대화하며, 상대 동의를 거쳐 실제 협업으로 이어지는 AI-native 커뮤니티 플랫폼입니다.",
        fill=PALE_TEAL,
        accent=TEAL,
    )
    add_heading(doc, "1. 왜 만드는가", 1)
    add_text(
        doc,
        "SparkClaw 참여팀은 서로의 제품명은 알아도 실제 해결 Task, 팀 경쟁력, 검증 근거와 연결 필요를 파악하기 어렵습니다. Arena는 프로그램 운영 정보가 아니라 '누구를 만나고 무엇을 함께할 것인가'를 해결합니다.",
        size=10.7,
        after=4,
        line=1.18,
    )
    for item in [
        "창업팀: 회계·재무, 고객 획득, PoC, 채용, 개발 외주와 같은 실제 문제를 동료에게 묻고 싶다.",
        "산업 파트너: 전체 참가기업 중 자사의 Task에 맞는 팀을 근거 중심으로 찾고 싶다.",
        "SparkLabs: 추천·소개·공지·활동 이력을 관리하면서 개인정보와 미확인 정보를 보호해야 한다.",
    ]:
        compact_bullet(doc, item)
    add_heading(doc, "2. 기존 프로그램 사이트와의 구분", 1)
    compact_bullet(doc, "프로그램 사이트: Weekly Report, 일정, 혜택, 과제와 내부 운영 정보.")
    compact_bullet(doc, "AI Arena: 기업 탐색, Community, 협업 요청, 동의 기반 소개와 개인 활동 기록.")
    add_callout(
        doc,
        "제품 판단 기준",
        "기능 수보다 '혼자서는 만나기 어려운 사람·기업·자원·기회를 실제로 만날 수 있는가'를 우선합니다.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    page_break(doc)

    # PAGE 2 - audiences
    add_heading(doc, "3. 핵심 사용자와 기대 가치", 1)
    role(
        doc,
        "SparkClaw 참여기업",
        "동료·파트너와 실제 협업 기회를 찾는 사용자",
        [
            "다른 팀의 서비스, 해결 Task, 팀 경쟁력과 정량 근거를 빠르게 이해한다.",
            "질문·경험·출시 소식을 나누고, 협업 요청과 응답을 My Log에서 확인한다.",
        ],
    )
    role(
        doc,
        "산업 파트너·외부 기업",
        "자사의 문제를 해결할 AI 스타트업을 찾는 사용자",
        [
            "산업명보다 구체적 Task와 공개 근거를 기준으로 후보를 탐색·비교한다.",
            "PoC·평가·API·채용·사업 제휴 목적을 전달하고 상대 동의 후 소개받는다.",
        ],
    )
    role(
        doc,
        "멘토·투자자·LP",
        "실행력과 성장 가능성을 판단하는 네트워크 사용자",
        [
            "팀 경력·학력, 문제, 최근 성과, 숫자, 검증 단계와 부족한 요소를 함께 본다.",
            "멘토링·투자 검토·산업 연결이 필요한 팀을 근거 중심으로 선별한다.",
        ],
    )
    role(
        doc,
        "SparkLabs 운영진",
        "큐레이션·소개·운영 품질을 책임지는 관리자",
        [
            "기업 프로필, Spotlight, 공식 공지와 추천 근거를 편집하고 품질을 검수한다.",
            "협업 요청의 발신·수신·승인·거절을 감사하고 비공개 정보와 AI 추정을 통제한다.",
        ],
    )
    add_callout(
        doc,
        "공통 원칙",
        "공개 가능한 근거만 보여주고, 연락처는 자동 공개하지 않으며, 소개는 상대 팀 승인 후 SparkLabs가 다음 단계를 조율합니다.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    page_break(doc)

    # PAGE 3 - product experiences
    add_heading(doc, "4. 핵심 제품 경험", 1)
    add_heading(doc, "Discover", 2)
    compact_bullet(doc, "Editorial Spotlight·What's New로 최근 실행 성과와 운영진 공지를 가장 먼저 보여준다.")
    compact_bullet(doc, "Spark AI 기업 탐색, Task-driven Search, Company Directory, Compare로 공개 근거를 탐색한다.")
    compact_bullet(doc, "협업 적합 기업·나와 비슷한 팀은 점수 대신 추천 순위와 기업별로 다른 선정 이유를 제공한다.")

    add_heading(doc, "Community", 2)
    compact_bullet(doc, "내용을 먼저 쓰면 Spark AI가 제목·채널·공개 범위를 제안하며 사용자가 수정할 수 있다.")
    compact_bullet(doc, "Pre-OT 공통 수요와 공식 공지에서 대화를 시작하고 글·댓글·upvote·수정 기능을 제공한다.")
    compact_bullet(doc, "Public은 산업 파트너 포함, Private은 부트캠프 멤버와 SparkLabs만 접근한다.")

    add_heading(doc, "동의 기반 협업", 2)
    add_lifecycle_table(doc)

    add_heading(doc, "Clawee 클로이와 My Log", 2)
    compact_bullet(doc, "Clawee는 실제 메뉴와 섹션을 이동하며 안내하고, 자연어 명령을 다음 행동으로 연결한다.")
    compact_bullet(doc, "My Log는 협업 요청·Community 활동·Bounty 상태를 최신순 raw log로 보여준다.")

    page_break(doc)

    # PAGE 4 - scope, bounty, principles
    add_heading(doc, "5. Bounty의 현재 위치", 1)
    add_callout(
        doc,
        "Release Gate",
        "실제 스폰서 문제와 운영 승인이 없는 데모 과제는 공개하지 않습니다. 현재는 준비 단계이며, 승인된 Sponsor Brief만 Bounty Board에 공개합니다.",
        fill="FFF7E6",
        accent=GOLD,
    )
    compact_bullet(doc, "Sponsor Brief 제출 → SparkLabs 검토 → 승인된 Bounty 공개 → 제출·검증 → PoC·투자·크레딧 연결")

    add_heading(doc, "6. 현재 반드시 포함할 범위", 1)
    must = [
        "기업 프로필과 Task 기반 Directory/Search",
        "Editorial Spotlight와 공식 공지",
        "기업 비교·유사 팀·협업 적합 기업 추천",
        "Community 글·댓글·upvote·수정·공개 범위",
        "상대 동의 기반 협업 요청과 운영 감사 로그",
        "My Log와 Supabase sc_arena_ 활동 원장",
        "역할별 권한 및 개인정보 보호",
        "Clawee 튜토리얼과 AI Arena 안내",
    ]
    for item in must:
        compact_bullet(doc, item)

    add_heading(doc, "7. 다음 단계", 1)
    for item in [
        "실제 승인된 Sponsor Brief 기반 Bounty 공개",
        "협업·댓글·Bounty·공식 공지의 인앱 및 이메일 알림",
        "멘토·투자사·LP 네트워크와 동의 기반 연결",
        "Clawee의 사용자 맥락 기억과 후속 제안",
    ]:
        compact_bullet(doc, item)

    add_heading(doc, "8. 운영 원칙과 성공 기준", 1)
    compact_bullet(doc, "Community first: 관계와 대화가 반복되는 구조를 우선한다.")
    compact_bullet(doc, "AI는 요약·추천하고 사람은 근거와 공개 범위를 검증한다.")
    compact_bullet(doc, "경쟁 점수보다 서로 다른 강점과 상호보완 가능성을 보여준다.")
    add_callout(
        doc,
        "최종 성공 기준",
        "발견한 팀이 대화하고 서로 동의한 뒤 실제 협업·PoC·멘토링·투자 검토로 이어지는가?",
        fill=PALE_TEAL,
        accent=TEAL,
    )
    add_text(
        doc,
        "정리 근거: 260806 AI Arena 피드백, 260811 수정 논의 미팅, 260812 호민 대표님 논의 및 2026년 8월 기준 구현·운영 기록",
        size=8.8,
        color=MUTED,
        italic=True,
        before=3,
        after=0,
    )

    doc.core_properties.title = "SparkClaw AI Arena 제품 방향 및 핵심 스펙 - 4쪽 압축본"
    doc.core_properties.subject = "Purpose, Audience & Core Product Scope"
    doc.core_properties.author = "SparkLabs"
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build())
